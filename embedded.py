'''# __define-ocg__

import os
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
import pickle

# =========================
# Load Embedding Model
# =========================

print("Loading model...")

model = SentenceTransformer(
    'all-MiniLM-L6-v2'
)

print("Model loaded!")

# =========================
# Document Folder
# =========================

DOCUMENT_FOLDER = "sample"

# =========================
# Read PDF
# =========================

def read_pdf(path):

    text = ""

    reader = PdfReader(path)

    for page in reader.pages:

        extracted = page.extract_text()

        if extracted:
            text += extracted + " "

    return text

# =========================
# Read DOCX
# =========================

def read_docx(path):

    doc = Document(path)

    return "\n".join(
        [p.text for p in doc.paragraphs]
    )

# =========================
# Read TXT
# =========================

def read_txt(path):

    with open(
        path,
        "r",
        encoding="utf-8"
    ) as f:

        return f.read()

# =========================
# Chunk Text
# =========================

def chunk_text(
    text,
    chunk_size=300
):

    words = text.split()

    chunks = []

    for i in range(
        0,
        len(words),
        chunk_size
    ):

        chunk = " ".join(
            words[i:i + chunk_size]
        )

        chunks.append(chunk)

    return chunks

# =========================
# Process Documents
# =========================

all_embeddings = []

print("Processing documents...")

for file in os.listdir(DOCUMENT_FOLDER):

    path = os.path.join(
        DOCUMENT_FOLDER,
        file
    )

    content = ""

    # PDF
    if file.endswith(".pdf"):

        content = read_pdf(path)

    # DOCX
    elif file.endswith(".docx"):

        content = read_docx(path)

    # TXT
    elif file.endswith(".txt"):

        content = read_txt(path)

    else:
        continue

    print(f"\nProcessing: {file}")

    # Split into chunks
    chunks = chunk_text(content)

    print(f"Chunks created: {len(chunks)}")

    # Create embeddings
    for chunk in chunks:

        embedding = model.encode(chunk)

        all_embeddings.append({

            "file": file,
            "text": chunk,
            "embedding": embedding

        })

    print("Embeddings created!")

print("\nDONE!")
print(f"Total embeddings: {len(all_embeddings)}")'''






# __define-ocg__ FULL EMBEDDING + FAISS INDEX CODE

import os
import pickle
import numpy as np
import faiss
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from charset_normalizer import from_path
# =========================
# Load Embedding Model
# =========================

print("Loading model...")

model = SentenceTransformer(
    'all-MiniLM-L6-v2'
)

print("Model loaded!")

# =========================
# Document Folder
# =========================

DOCUMENT_FOLDER = "sample"

# =========================
# Store All Data
# =========================

all_embeddings = []

# =========================
# Read PDF Function
# =========================

def read_pdf(path):

    text = ""
    try:
        
        reader = PdfReader(path)

        for page in reader.pages:

            extracted = page.extract_text()

            if extracted:

                text += extracted + " "
    except Exception as e:
        print(f"[ERROR] PDF failed: {path}")
        print(e)
    return text

# =========================
# Read DOCX Function
# =========================

def read_docx(path):
    try:
        doc = Document(path)
        return "\n".join(
            [p.text for p in doc.paragraphs]
        )
    except Exception as e:
        print(f"[ERROR] DOCX failed: {path}")
        print(e)
        return ""

# =========================
# Read TXT Function
# =========================

def read_txt(path):
    try:
        result = from_path(path).best()
        encoding =(
            result.encoding
            if result
            else "utf-8"     
        )
        with open(
            path,
            "r",
            encoding="utf-8"
         ) as f:

              return f.read()

    except Exception as e:
        print(f"[ERROR] TXT failed: {path}")
        print(e)
        return ""


# =========================
# Read XLSX
# =========================
def read_xlsx(path):
    try:
        df = pd.read_excel(path, engine = "openpyxl")
        return df.astype(str).to_string()
    except Exception as e:
        print(f"[ERROR] XLSX failed: {path}")
        print(e)
        return ""

# =========================
# Chunk Text Function
# =========================

def chunk_text(
    text,
    chunk_size=300
):

    words = text.split()

    chunks = []

    for i in range(
        0,
        len(words),
        chunk_size
    ):

        chunk = " ".join(

            words[i:i + chunk_size]

        )

        chunks.append(chunk)

    return chunks

# =========================
# Process Documents
# =========================

print("\nProcessing documents...")

for file in os.listdir(DOCUMENT_FOLDER):

    path = os.path.join(
        DOCUMENT_FOLDER,
        file
    )
    if not os.path.isfile(path):
        continue

    content = ""
    print(f"\nProcessing: {file}")
    # =========================
    # PDF
    # =========================

    if file.endswith(".pdf"):

        content = read_pdf(path)

    # =========================
    # DOCX
    # =========================

    elif file.endswith(".docx"):

        content = read_docx(path)

    # =========================
    # TXT
    # =========================

    elif file.endswith(".txt"):

        content = read_txt(path)
    # =========================
    # XLSX
    # =========================
    elif file.endswith(".xlsx"):
        content = read_xlsx(path)
    else:
        print("[SKIPPED] Unsupported file")
        continue

    
    # =========================
    # Empty check
    # =========================
    if not content.strip():
        print("[WARNING] Empty content")
        continue
    # =========================
    # Create Chunks
    # =========================

    chunks = chunk_text(content)

    print(f"Chunks created: {len(chunks)}")
    # =========================
    # Batch Embedding
    # =========================
    embedding = model.encode(
        chunks, batch_size =32, show_progress_bar=True
    )
    for chunk, embedding in zip(chunks, embedding):
        all_embeddings.append({
            "file":file,
            "text":chunk,
            "embedding":embedding
         })
    print("Embeddings created!")
     
# =========================
# NO DOCUMENT CHECK
# =========================
if len(all_embeddings) == 0:
    print("\nNo embeddings created!")
    exit(0)

# =========================
# Convert Embeddings to NumPy
# =========================

embedding_vectors = np.array(
    [item["embedding"]
     for item in all_embeddings]

).astype("float32")

print("\nEmbedding matrix shape:")

print(embedding_vectors.shape)

# =========================
# Create FAISS Index
# =========================

dimension = embedding_vectors.shape[1]

index = faiss.IndexFlatL2(
    dimension
)

index.add(
    embedding_vectors
)

print("\nFAISS index created!")

# =========================
# Save FAISS Index
# =========================

faiss.write_index(

    index,

    "document_index.faiss"

)

print("FAISS index saved!")

# =========================
# Save Metadata
# =========================

with open(
    "metadata.pkl",
    "wb"
) as f:

    pickle.dump(
        all_embeddings,
        f
    )

print("Metadata saved!")

# =========================
# DONE
# =========================

print("\nDONE SUCCESSFULLY!")
